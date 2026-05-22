import argparse
import hashlib
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import yaml
from tqdm import tqdm

logger = logging.getLogger(__name__)

_ocr_instance = None


def get_ocr():
    global _ocr_instance
    if _ocr_instance is None:
        logger.info("Initializing PaddleOCR (first time only)...")
        from paddleocr import PaddleOCR
        _ocr_instance = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)
        logger.info("PaddleOCR initialized successfully")
    return _ocr_instance


def sanitize_filename(filename: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    sanitized = sanitized.strip('. ')
    return sanitized[:200] if sanitized else "untitled"


def resolve_output_path(output_dir: Path, stem: str, source_path: Path) -> Path:
    safe_stem = sanitize_filename(stem)
    out_path = output_dir / f"{safe_stem}.md"

    if not out_path.exists():
        return out_path

    source_hash = hashlib.md5(str(source_path).encode()).hexdigest()[:8]
    return output_dir / f"{safe_stem}_{source_hash}.md"


def parse_pdf(file_path: str) -> dict[str, Any]:
    import pdfplumber

    texts = []
    tables = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                texts.append(f"<!-- Page {i+1} -->\n{text}")

            page_tables = page.extract_tables()
            for table in page_tables:
                if table:
                    md_table = format_table_as_markdown(table)
                    tables.append(md_table)

    return {
        "title": Path(file_path).stem,
        "content": "\n\n".join(texts),
        "tables": tables,
    }


def parse_docx(file_path: str) -> dict[str, Any]:
    import docx

    doc = docx.Document(file_path)
    texts = []
    tables = []

    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(format_table_as_markdown(rows))

    return {
        "title": Path(file_path).stem,
        "content": "\n\n".join(texts),
        "tables": tables,
    }


def parse_image(file_path: str) -> dict[str, Any]:
    ocr = get_ocr()
    result = ocr.ocr(file_path, cls=True)

    texts = []
    for line in result:
        if line:
            for word_info in line:
                if word_info and len(word_info) >= 2 and word_info[1]:
                    texts.append(str(word_info[1][0]))

    return {
        "title": Path(file_path).stem,
        "content": "\n".join(texts),
        "tables": [],
    }


def escape_markdown_cell(text: str) -> str:
    normalized = " ".join(text.splitlines())
    return normalized.replace("|", "\\|").strip()


def format_table_as_markdown(table: list[list[Any]]) -> str:
    if not table:
        return ""

    max_cols = max(len(row) for row in table) if table else 0
    if max_cols == 0:
        return ""

    rows = []
    for row in table:
        cells = [escape_markdown_cell(str(cell)) if cell else "" for cell in row]
        cells.extend([""] * (max_cols - len(cells)))
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) < 1:
        return ""

    separator = "| " + " | ".join(["---"] * max_cols) + " |"

    return "\n".join([rows[0], separator] + rows[1:])


def build_markdown(title: str, content: str, tables: list[str], source_file: str) -> str:
    safe_title = " ".join(title.splitlines()).strip()

    frontmatter = {
        "标题": safe_title,
        "来源": source_file,
        "日期": datetime.now().strftime("%Y-%m-%d"),
        "解析工具": "deepdoc-toolkit",
    }

    frontmatter_str = yaml.dump(frontmatter, allow_unicode=True, default_flow_style=False).strip()

    parts = [f"---\n{frontmatter_str}\n---\n"]

    if content.strip():
        parts.append(content.strip())

    if tables:
        parts.append("\n\n## 表格\n")
        for i, table in enumerate(tables, 1):
            parts.append(f"\n### 表格 {i}\n\n{table}\n")

    return "\n\n".join(parts)


def parse_single_file(file_path: str, output_dir: Optional[str] = None) -> Optional[str]:
    path = Path(file_path)

    if not path.exists():
        logger.error(f"File not found: {file_path}")
        return None

    if not path.is_file():
        logger.error(f"Not a file: {file_path}")
        return None

    ext = path.suffix.lower()

    if ext == ".pdf":
        result = parse_pdf(file_path)
    elif ext == ".docx":
        result = parse_docx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        result = parse_image(file_path)
    else:
        logger.warning(f"Unsupported format: {ext}")
        return None

    md_content = build_markdown(
        title=result["title"],
        content=result["content"],
        tables=result["tables"],
        source_file=path.name,
    )

    if output_dir:
        out_path = resolve_output_path(Path(output_dir), path.stem, path)
    else:
        out_path = resolve_output_path(path.parent, path.stem, path)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(md_content, encoding="utf-8")

    return str(out_path)


def batch_parse(input_dir: str, output_dir: str) -> None:
    input_path = Path(input_dir)
    output_path = Path(output_dir)

    if input_path.resolve() == output_path.resolve():
        logger.error("Input and output directories must be different")
        return

    output_path.mkdir(parents=True, exist_ok=True)

    supported = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".tiff", ".bmp"}
    files = [f for f in input_path.rglob("*") if f.suffix.lower() in supported]

    if not files:
        logger.warning(f"No supported files found in {input_dir}")
        return

    logger.info(f"Found {len(files)} files to parse")

    success_count = 0
    error_count = 0

    for file in tqdm(files, desc="Parsing files"):
        try:
            result = parse_single_file(str(file), str(output_path))
            if result:
                success_count += 1
                logger.debug(f"  -> {result}")
        except Exception as e:
            error_count += 1
            logger.error(f"  -> Error parsing {file.name}: {e}")

    logger.info(f"Parsing complete: {success_count} succeeded, {error_count} failed")


def main() -> None:
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
    )

    parser = argparse.ArgumentParser(description="DeepDoc Toolkit - Document parser for game-kb")
    parser.add_argument("input", nargs="?", help="Single file to parse")
    parser.add_argument("--input", dest="input_dir", help="Input directory for batch parsing")
    parser.add_argument("--output", dest="output_dir", default="./output", help="Output directory (default: ./output)")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose logging")

    args = parser.parse_args()

    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)

    if args.input_dir:
        batch_parse(args.input_dir, args.output_dir)
    elif args.input:
        result = parse_single_file(args.input, args.output_dir)
        if result:
            logger.info(f"Output: {result}")
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
